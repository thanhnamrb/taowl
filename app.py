from __future__ import annotations

import csv
import hashlib
import io
import re
import tempfile
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from zipfile import ZipFile

import streamlit as st
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

# ===================== DATA MODEL =====================
@dataclass(slots=True)
class VocabWord:
    word: str
    word_type: str = ""
    pronunciation: str = ""
    meaning: str = ""


@dataclass(slots=True)
class VocabFamily:
    number: str
    words: list[VocabWord] = field(default_factory=list)


@dataclass(slots=True)
class VocabDocument:
    unit: str
    title: str
    document_type: str = "VOCAB BUILDER"
    section_label: str = "UNIT"
    heading_unit: str = ""
    families: list[VocabFamily] = field(default_factory=list)

    @property
    def word_count(self) -> int:
        return sum(len(f.words) for f in self.families)

    @property
    def family_count(self) -> int:
        return len(self.families)

    @property
    def unit_badge(self) -> str:
        # The user's corrected reference uses the full sub-unit in the badge: 5.1.
        return (self.unit or "").strip() or "0"

    @property
    def display_heading_unit(self) -> str:
        # 5.1 -> UNIT 5 by default, matching the corrected reference.
        if self.heading_unit.strip():
            return self.heading_unit.strip()
        return (self.unit or "").split(".", 1)[0].strip() or self.unit

# ===================== PARSER =====================
HEADER_ALIASES = {
    "no": {"no", "no.", "number", "stt", "stt."},
    "word": {"word", "vocabulary", "vocab"},
    "type": {"type", "word type", "pos", "part of speech"},
}


def _looks_like_header(row: list[str]) -> bool:
    cells = [c.strip().lower() for c in row]
    if len(cells) < 3:
        return False
    return (
        cells[0] in HEADER_ALIASES["no"]
        and cells[1] in HEADER_ALIASES["word"]
        and cells[2] in HEADER_ALIASES["type"]
    )


def parse_vocab_csv(
    raw_data: str,
    *,
    unit: str,
    title: str,
    document_type: str = "VOCAB BUILDER",
    section_label: str = "UNIT",
    heading_unit: str = "",
) -> VocabDocument:
    """Parse CSV/clipboard data into a semantic document model.

    Expected columns: No., Word, Type, Pronunciation, Meaning.
    A blank No. means the row belongs to the previous word family.
    """
    stream = io.StringIO((raw_data or "").strip())
    rows = list(csv.reader(stream))
    rows = [r for r in rows if r and "".join(r).strip()]

    if rows and _looks_like_header(rows[0]):
        rows = rows[1:]

    document = VocabDocument(
        unit=(unit or "").strip(),
        title=(title or "").strip(),
        document_type=(document_type or "VOCAB BUILDER").strip().upper(),
        section_label=(section_label or "UNIT").strip().upper(),
        heading_unit=(heading_unit or "").strip(),
    )

    current_family: VocabFamily | None = None

    for source_index, row in enumerate(rows, start=1):
        row = list(row)
        while len(row) < 5:
            row.append("")
        row = row[:5]

        no, word, word_type, pronunciation, meaning = (c.strip() for c in row)

        if not word:
            # The validator will catch meaningful malformed rows. Pure blank rows were removed above.
            continue

        if no:
            current_family = VocabFamily(number=no)
            document.families.append(current_family)
        elif current_family is None:
            # Preserve malformed input as a synthetic family so validation can explain it.
            current_family = VocabFamily(number="")
            document.families.append(current_family)

        current_family.words.append(
            VocabWord(
                word=word,
                word_type=word_type,
                pronunciation=pronunciation,
                meaning=meaning,
            )
        )

    return document

# ===================== THEME =====================
@dataclass(slots=True)
class TextStyle:
    font: str = "Arial"
    size: float = 10.5
    color: str = "000000"
    bold: bool = False
    italic: bool = False


@dataclass(slots=True)
class HeaderSegment:
    line: int
    text: str
    style: TextStyle


@dataclass(slots=True)
class TableTheme:
    # Geometry/spacing remains from the original pre-app VOCAB table.
    # Branding is applied without changing the row rhythm.
    no_fill: str = "BF4E14"
    header_fill: str = "0F4761"
    no_header_text: TextStyle = field(
        default_factory=lambda: TextStyle("Orbitron", 10.0, "FFFFFF", True, False)
    )
    header_text: TextStyle = field(
        default_factory=lambda: TextStyle("Arial", 10.5, "FFFFFF", True, False)
    )
    no_body_text: TextStyle = field(
        default_factory=lambda: TextStyle("Orbitron", 10.0, "FFFFFF", True, False)
    )
    body_text: TextStyle = field(
        default_factory=lambda: TextStyle("Times New Roman", 12, "000000", False, False)
    )
    outer_border_color: str = "0F4761"
    outer_border_pt: float = 1.0
    inner_border_color: str = "B8C8D0"
    inner_border_pt: float = 0.5


@dataclass(slots=True)
class FooterTheme:
    slogan_segments: list[HeaderSegment] = field(default_factory=list)
    slogan_align: str = "left"
    phone: str = "0345286842"
    email: str = "email@yourcenter.com"
    phone_symbol: str = "☎"
    email_symbol: str = "✉"
    contact_font: str = "Arial"
    contact_size: float = 8.0
    phone_color: str = "BF4E14"
    email_color: str = "0F4761"
    contact_bold: bool = True
    page_fill: str = "BF4E14"
    page_text: TextStyle = field(
        default_factory=lambda: TextStyle("Orbitron", 9.0, "FFFFFF", True, False)
    )
    page_badge_size_pt: float = 24.0


@dataclass(slots=True)
class ThemeConfig:
    badge: TextStyle = field(
        default_factory=lambda: TextStyle("Orbitron", 48, "FFFFFF", True)
    )
    document_type: TextStyle = field(
        default_factory=lambda: TextStyle("Monument Extended", 18, "BF4E14", False)
    )
    heading: TextStyle = field(
        default_factory=lambda: TextStyle("Montserrat", 27, "000000", True)
    )
    header_segments: list[HeaderSegment] = field(default_factory=list)
    header_alignments: dict[int, str] = field(
        default_factory=lambda: {1: "right", 2: "right", 3: "right"}
    )
    # Extra body clearance under the header. Renderer automatically adds a little more
    # if a third header line is used.
    title_clearance_pt: float = 28.0
    table: TableTheme = field(default_factory=TableTheme)
    footer: FooterTheme = field(default_factory=FooterTheme)


def default_header_segments() -> list[HeaderSegment]:
    # Default is the user's second reference image.
    peach = "E7A07F"
    blue = "7FA9BC"
    return [
        HeaderSegment(1, "Supplementary", TextStyle("Arial", 10.5, peach, False)),
        HeaderSegment(1, "Vocabulary", TextStyle("Arial", 10.5, peach, False)),
        HeaderSegment(2, "for", TextStyle("Arial", 10.5, peach, False)),
        HeaderSegment(2, "KET", TextStyle("Orbitron", 10.5, blue, True)),
        HeaderSegment(2, "Learners", TextStyle("Orbitron", 10.5, blue, True)),
        HeaderSegment(2, "in", TextStyle("Arial", 10.5, peach, False)),
        HeaderSegment(2, "Orwell", TextStyle("Orbitron", 10.5, blue, True)),
        HeaderSegment(2, "Classes", TextStyle("Orbitron", 10.5, blue, True)),
    ]


def default_footer_segments() -> list[HeaderSegment]:
    return [
        HeaderSegment(1, "From", TextStyle("Arial", 9, "BF4E14", False)),
        HeaderSegment(1, "Learners", TextStyle("Orbitron", 9, "0F4761", True)),
        HeaderSegment(1, "to", TextStyle("Arial", 9, "BF4E14", False)),
        HeaderSegment(1, "Explorers", TextStyle("Orbitron", 9, "0F4761", True)),
    ]


def default_theme() -> ThemeConfig:
    return ThemeConfig(
        header_segments=default_header_segments(),
        footer=FooterTheme(slogan_segments=default_footer_segments()),
    )

# ===================== RENDERER =====================
def _hex(value: str, fallback: str = "000000") -> str:
    v = (value or "").strip().lstrip("#").upper()
    return v if len(v) == 6 and all(ch in "0123456789ABCDEF" for ch in v) else fallback


def _apply_run_style(run, style: TextStyle) -> None:
    run.font.name = style.font
    run.font.size = Pt(float(style.size))
    run.font.bold = bool(style.bold)
    run.font.italic = bool(style.italic)
    run.font.color.rgb = RGBColor.from_string(_hex(style.color))
    rfonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), style.font)


def _remove_text_runs_preserve_drawings(paragraph) -> None:
    for run in list(paragraph.runs):
        if not run._r.xpath(".//w:drawing | .//w:pict"):
            run._element.getparent().remove(run._element)


def _set_single_run_paragraph(paragraph, text: str, style: TextStyle, *, align=None) -> None:
    _remove_text_runs_preserve_drawings(paragraph)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    _apply_run_style(run, style)


def _set_title_block(doc: Document, document: VocabDocument, theme: ThemeConfig) -> None:
    title = doc.tables[0]

    badge = title.rows[0].cells[0]
    badge.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_single_run_paragraph(
        badge.paragraphs[0], document.unit_badge, theme.badge,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    type_cell = title.rows[0].cells[2]
    type_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_single_run_paragraph(
        type_cell.paragraphs[0], document.document_type.upper(), theme.document_type,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    main = title.rows[1].cells[2]
    main.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    while len(main.paragraphs) > 1:
        p = main.paragraphs[-1]._element
        p.getparent().remove(p)
    heading = f"{document.section_label} {document.display_heading_unit}: {document.title.upper()}"
    _set_single_run_paragraph(main.paragraphs[0], heading, theme.heading, align=WD_ALIGN_PARAGRAPH.LEFT)


def _alignment(value: str):
    value = (value or "center").lower()
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(value, WD_ALIGN_PARAGRAPH.CENTER)


def _set_header(doc: Document, theme: ThemeConfig) -> None:
    """Render 1-3 independently styled header lines.

    The logo drawing stays intact. Each HeaderSegment can be a word or a whole phrase,
    so the user can add/remove/reorder wording without changing Python.
    """
    header = doc.sections[0].header
    used_lines = max([s.line for s in theme.header_segments if s.text.strip()] or [1])
    used_lines = max(1, min(3, used_lines))

    while len(header.paragraphs) < used_lines:
        header.add_paragraph()

    # Clear text on the first 3 paragraphs but preserve floating drawings/logo.
    for line_no in range(1, 4):
        if len(header.paragraphs) < line_no:
            break
        p = header.paragraphs[line_no - 1]
        _remove_text_runs_preserve_drawings(p)
        p.alignment = _alignment(theme.header_alignments.get(line_no, "center"))
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        items = [s for s in theme.header_segments if int(s.line) == line_no and s.text.strip()]
        for index, segment in enumerate(items):
            # A segment may be a word OR a phrase. This makes header editing flexible.
            text = segment.text.strip()
            if index < len(items) - 1:
                text += " "
            run = p.add_run(text)
            _apply_run_style(run, segment.style)

    # Hide any unused extra header text paragraphs while keeping drawing runs.
    for p in header.paragraphs[used_lines:3]:
        _remove_text_runs_preserve_drawings(p)


def _set_title_clearance(doc: Document, theme: ThemeConfig) -> None:
    """Prevent title artwork from entering the header zone.

    The V3 template contains a blank paragraph directly before the title table.
    We control its spacing here. A third header line automatically receives more room.
    """
    title_tbl = doc.tables[0]._tbl
    prev = title_tbl.getprevious()
    if prev is None or prev.tag != qn("w:p"):
        p = OxmlElement("w:p")
        title_tbl.addprevious(p)
        prev = p

    used_lines = max([s.line for s in theme.header_segments if s.text.strip()] or [1])
    extra = max(0, min(3, used_lines) - 2) * 14.0
    clearance = max(18.0, float(theme.title_clearance_pt) + extra)

    ppr = prev.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    # Keep this paragraph visually empty and use only after-spacing as clearance.
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), str(int(clearance * 20)))  # pt -> twips
    spacing.set(qn("w:line"), "20")
    spacing.set(qn("w:lineRule"), "exact")


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        tr_pr.append(node)
    node.set(qn("w:val"), "true")


def _set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:cantSplit"))
    if node is None:
        node = OxmlElement("w:cantSplit")
        tr_pr.append(node)


def _set_family_keep_next(row, enabled: bool) -> None:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.keep_with_next = enabled
            p.paragraph_format.keep_together = True


def _clear_vmerge(cell) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    vm = tcpr.find(qn("w:vMerge"))
    if vm is not None:
        tcpr.remove(vm)


def _set_vmerge_tc(tc, *, restart: bool) -> None:
    tcpr = tc.get_or_add_tcPr()
    vm = tcpr.find(qn("w:vMerge"))
    if vm is None:
        vm = OxmlElement("w:vMerge")
        tcpr.append(vm)
    if restart:
        vm.set(qn("w:val"), "restart")
    else:
        vm.attrib.pop(qn("w:val"), None)




def _set_cell_no_wrap(cell, enabled: bool = True) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    node = tcpr.find(qn("w:noWrap"))
    if enabled and node is None:
        node = OxmlElement("w:noWrap")
        tcpr.append(node)
    elif not enabled and node is not None:
        tcpr.remove(node)



def _set_fixed_table_layout_and_no_column(table, min_no_inches: float = 0.50) -> None:
    """Force a fixed table grid and keep the No. column wide enough for one line.

    The original VOCAB geometry is preserved as much as possible: only the small
    amount needed by the No. column is borrowed from the Meaning column.
    """
    tblpr = table._tbl.tblPr
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid_cols = table._tbl.tblGrid.gridCol_lst
    if len(grid_cols) < 5:
        return

    widths = [int(col.w) for col in grid_cols]
    min_no_emu = int(Inches(float(min_no_inches)))
    if widths[0] < min_no_emu:
        delta = min_no_emu - widths[0]
        widths[0] = min_no_emu
        # Borrow only from Meaning so Word/Type/Pronunciation keep the original feel.
        widths[-1] = max(int(Inches(1.50)), widths[-1] - delta)

    for col, width in zip(grid_cols, widths):
        col.w = width

    # tcW in dxa/twips makes Word respect the same fixed geometry in every row.
    for row in table.rows:
        for i, cell in enumerate(row.cells[:len(widths)]):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(max(1, round(widths[i] / 635))))
        if row.cells:
            _set_cell_no_wrap(row.cells[0], True)


def _set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(fill, "FFFFFF"))


def _set_table_borders(table, theme: TableTheme) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)

    def set_edge(name: str, color: str, pt: float):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(max(1, int(round(float(pt) * 8)))))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), _hex(color, "0F4761"))

    for name in ("top", "left", "bottom", "right"):
        set_edge(name, theme.outer_border_color, theme.outer_border_pt)
    for name in ("insideH", "insideV"):
        set_edge(name, theme.inner_border_color, theme.inner_border_pt)


def _set_cell_text_preserve_geometry(
    cell, text: str, *, style: TextStyle, center: bool = False
) -> None:
    """Replace text while leaving the original cell margins/row geometry untouched."""
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    if p.runs:
        p.runs[0].text = text or ""
        _apply_run_style(p.runs[0], style)
        for r in list(p.runs[1:]):
            r._element.getparent().remove(r._element)
    else:
        r = p.add_run(text or "")
        _apply_run_style(r, style)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _style_header_row(row, table_theme: TableTheme) -> None:
    labels = ["No.", "Word", "Type", "Pronunciation", "Meaning"]
    for i, (cell, label) in enumerate(zip(row.cells, labels)):
        _set_cell_shading(cell, table_theme.no_fill if i == 0 else table_theme.header_fill)
        _set_cell_no_wrap(cell, True)
        _set_cell_text_preserve_geometry(
            cell, label,
            style=table_theme.no_header_text if i == 0 else table_theme.header_text,
            center=True,
        )


def _populate_vocab_table(doc: Document, document: VocabDocument, table_theme: TableTheme) -> None:
    if len(doc.tables) < 2:
        raise RuntimeError("Master template phải có title table và VOCAB prototype table.")
    table = doc.tables[1]
    if len(table.rows) < 3:
        raise RuntimeError("VOCAB prototype table cần header + 2 prototype rows.")

    table.autofit = False
    _set_fixed_table_layout_and_no_column(table, 0.50)
    header = table.rows[0]
    _set_repeat_header(header)
    _set_cant_split(header)
    _style_header_row(header, table_theme)
    _set_table_borders(table, table_theme)

    root_proto = table.rows[1]
    derivative_proto = table.rows[2]
    merge_groups = []

    for family in document.families:
        family_rows = []
        for word_index, item in enumerate(family.words):
            source = root_proto if word_index == 0 else derivative_proto
            table._tbl.append(deepcopy(source._tr))
            row = table.rows[-1]
            _set_cant_split(row)
            for cell_index, c in enumerate(row.cells):
                _clear_vmerge(c)
                _set_cell_shading(c, table_theme.no_fill if cell_index == 0 else "FFFFFF")

            # No. column is a branded orange rail; vocabulary text remains neutral/plain.
            _set_cell_text_preserve_geometry(
                row.cells[0], family.number if word_index == 0 else "",
                style=table_theme.no_body_text, center=True,
            )
            _set_cell_text_preserve_geometry(row.cells[1], item.word, style=table_theme.body_text)
            _set_cell_text_preserve_geometry(row.cells[2], item.word_type, style=table_theme.body_text)
            _set_cell_text_preserve_geometry(row.cells[3], item.pronunciation, style=table_theme.body_text)
            _set_cell_text_preserve_geometry(row.cells[4], item.meaning, style=table_theme.body_text)
            family_rows.append(row)

        for i, row in enumerate(family_rows):
            _set_family_keep_next(row, i < len(family_rows) - 1)
        if len(family_rows) > 1:
            merge_groups.append(family_rows)

    table._tbl.remove(root_proto._tr)
    table._tbl.remove(derivative_proto._tr)

    # One single continuous table. Apply vMerge only after all body text is final.
    for family_rows in merge_groups:
        _set_vmerge_tc(family_rows[0]._tr.tc_lst[0], restart=True)
        for row in family_rows[1:]:
            _set_vmerge_tc(row._tr.tc_lst[0], restart=False)



def _clear_footer(footer) -> None:
    for child in list(footer._element):
        footer._element.remove(child)


def _set_table_no_borders(table) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "nil")


def _append_page_field_run(parent, style: TextStyle, *, cached_text: str | None = None, field_type: str | None = None, instruction: str | None = None):
    rr = OxmlElement("w:r")
    rp = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for a in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{a}"), style.font)
    rp.append(fonts)
    if style.bold:
        rp.append(OxmlElement("w:b"))
    if style.italic:
        rp.append(OxmlElement("w:i"))
    col = OxmlElement("w:color")
    col.set(qn("w:val"), _hex(style.color, "FFFFFF"))
    rp.append(col)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(round(float(style.size) * 2))))
    rp.append(sz)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), str(int(round(float(style.size) * 2))))
    rp.append(szcs)
    rr.append(rp)
    if field_type:
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), field_type)
        rr.append(fld)
    elif instruction is not None:
        it = OxmlElement("w:instrText")
        it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        it.text = instruction
        rr.append(it)
    else:
        t = OxmlElement("w:t")
        t.text = cached_text or "1"
        rr.append(t)
    parent.append(rr)


def _add_page_badge(paragraph, footer_theme: FooterTheme) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    pict = OxmlElement("w:pict")
    oval = etree.Element("{urn:schemas-microsoft-com:vml}oval")
    size = max(18.0, min(40.0, float(footer_theme.page_badge_size_pt)))
    oval.set("style", f"width:{size}pt;height:{size}pt")
    oval.set("fillcolor", f"#{_hex(footer_theme.page_fill, 'BF4E14')}")
    oval.set("stroked", "f")
    textbox = etree.Element("{urn:schemas-microsoft-com:vml}textbox")
    textbox.set("inset", "0,0,0,0")
    txc = OxmlElement("w:txbxContent")
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    ppr.append(jc)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), str(int(max(10.0, footer_theme.page_text.size * 1.35) * 20)))
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    p.append(ppr)
    _append_page_field_run(p, footer_theme.page_text, field_type="begin")
    _append_page_field_run(p, footer_theme.page_text, instruction=" PAGE ")
    _append_page_field_run(p, footer_theme.page_text, field_type="separate")
    _append_page_field_run(p, footer_theme.page_text, cached_text="1")
    _append_page_field_run(p, footer_theme.page_text, field_type="end")
    txc.append(p)
    textbox.append(txc)
    oval.append(textbox)
    pict.append(oval)
    run._r.append(pict)


def _set_footer(doc: Document, footer_theme: FooterTheme) -> None:
    section = doc.sections[0]
    footer = section.footer
    _clear_footer(footer)
    content_width = section.page_width - section.left_margin - section.right_margin
    table = footer.add_table(rows=1, cols=3, width=content_width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # Balanced columns make the page badge mathematically centered in the text area.
    left_w = Inches(3.0)
    center_w = Inches(0.78)
    right_w = content_width - left_w - center_w
    widths = [left_w, center_w, right_w]
    for i, width in enumerate(widths):
        table.columns[i].width = width
        table.cell(0, i).width = width
    _set_table_no_borders(table)

    left = table.cell(0, 0)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]
    p.alignment = _alignment(footer_theme.slogan_align)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    items = [s for s in footer_theme.slogan_segments if s.text.strip()]
    for i, segment in enumerate(items):
        text = segment.text.strip() + (" " if i < len(items) - 1 else "")
        r = p.add_run(text)
        _apply_run_style(r, segment.style)

    center = table.cell(0, 1)
    center.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _add_page_badge(center.paragraphs[0], footer_theme)

    right = table.cell(0, 2)
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p1 = right.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r = p1.add_run(f"{footer_theme.phone_symbol} {footer_theme.phone}".strip())
    _apply_run_style(r, TextStyle(footer_theme.contact_font, footer_theme.contact_size, footer_theme.phone_color, footer_theme.contact_bold))
    p2 = right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(f"{footer_theme.email_symbol} {footer_theme.email}".strip())
    _apply_run_style(r, TextStyle(footer_theme.contact_font, footer_theme.contact_size, footer_theme.email_color, footer_theme.contact_bold))

    # Word requires a paragraph after a table in a header/footer; make it visually negligible.
    tail = footer.add_paragraph()
    tail.paragraph_format.space_before = Pt(0)
    tail.paragraph_format.space_after = Pt(0)
    tail.paragraph_format.line_spacing = Pt(1)


def render_vocab_docx(
    document: VocabDocument,
    *,
    template_path: str | Path,
    theme: ThemeConfig | None = None,
) -> io.BytesIO:
    theme = theme or default_theme()
    doc = Document(str(template_path))
    _set_header(doc, theme)
    _set_title_clearance(doc, theme)
    _set_title_block(doc, document, theme)
    _populate_vocab_table(doc, document, theme.table)
    _set_footer(doc, theme.footer)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ===================== VALIDATION =====================
@dataclass(slots=True)
class ValidationMessage:
    level: str  # error | warning
    message: str


def validate_vocab(document: VocabDocument) -> list[ValidationMessage]:
    messages: list[ValidationMessage] = []

    if not document.unit:
        messages.append(ValidationMessage("error", "Unit không được để trống."))
    if not document.title:
        messages.append(ValidationMessage("error", "Title không được để trống."))
    if not document.families:
        messages.append(ValidationMessage("error", "Không tìm thấy dữ liệu từ vựng hợp lệ."))
        return messages

    seen_numbers: set[str] = set()
    seen_words: dict[str, int] = {}

    for family_index, family in enumerate(document.families, start=1):
        if not family.number:
            messages.append(
                ValidationMessage(
                    "error",
                    f"Family #{family_index} không có STT. Dòng đầu tiên phải có giá trị ở cột No.",
                )
            )
        elif family.number in seen_numbers:
            messages.append(
                ValidationMessage("warning", f"STT {family.number} xuất hiện nhiều hơn một lần.")
            )
        seen_numbers.add(family.number)

        if not family.words:
            messages.append(ValidationMessage("error", f"Family {family.number} không có từ."))

        for item in family.words:
            key = item.word.casefold()
            seen_words[key] = seen_words.get(key, 0) + 1
            if not item.word_type:
                messages.append(
                    ValidationMessage(
                        "warning", f"Từ “{item.word}” chưa có Type/part of speech."
                    )
                )

    duplicates = sorted(word for word, count in seen_words.items() if count > 1)
    if duplicates:
        preview = ", ".join(duplicates[:8])
        suffix = "…" if len(duplicates) > 8 else ""
        messages.append(ValidationMessage("warning", f"Có từ lặp: {preview}{suffix}"))

    return messages


def has_errors(messages: list[ValidationMessage]) -> bool:
    return any(m.level == "error" for m in messages)

# ===================== QA =====================
@dataclass(slots=True)
class QAResult:
    ok: bool
    checks: list[str]
    problems: list[str]


def _docx_text_fallback(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_docx_markdown(path: str | Path) -> str:
    """Use Microsoft's MarkItDown when installed; fall back locally for resilience."""
    path = Path(path)
    try:
        from markitdown import MarkItDown  # type: ignore

        result = MarkItDown().convert(str(path))
        return result.text_content or ""
    except Exception:
        return _docx_text_fallback(path)


def run_content_and_brand_qa(path: str | Path, expected: VocabDocument) -> QAResult:
    path = Path(path)
    checks: list[str] = []
    problems: list[str] = []

    text = extract_docx_markdown(path)
    expected_title_tokens = [expected.document_type, f"{expected.section_label} {expected.display_heading_unit}", expected.title]
    for token in expected_title_tokens:
        if token.casefold() in text.casefold():
            checks.append(f"Title token OK: {token}")
        else:
            problems.append(f"Thiếu title token: {token}")

    for family in expected.families:
        for item in family.words:
            if item.word.casefold() not in text.casefold():
                problems.append(f"Thiếu từ sau khi render: {item.word}")
    if not any(p.startswith("Thiếu từ") for p in problems):
        checks.append(f"Đủ {expected.word_count} vocabulary rows về mặt semantic.")

    # OOXML brand checks: these are structural, not visual.
    with ZipFile(path) as zf:
        names = set(zf.namelist())
        document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        styles_xml = zf.read("word/styles.xml").decode("utf-8", errors="ignore")
        header_xml = "\n".join(
            zf.read(n).decode("utf-8", errors="ignore")
            for n in names
            if n.startswith("word/header") and n.endswith(".xml")
        )
        footer_xml = "\n".join(
            zf.read(n).decode("utf-8", errors="ignore")
            for n in names
            if n.startswith("word/footer") and n.endswith(".xml")
        )

    for token, label in [
        ("BF4E14", "orange BF4E14"),
        ("0F4761", "navy 0F4761"),
        ("Arial", "Arial"),
        ("Montserrat", "Montserrat"),
        ("Orbitron", "Orbitron"),
    ]:
        haystack = document_xml + styles_xml + header_xml + footer_xml
        if token in haystack:
            checks.append(f"Brand token OK: {label}")
        else:
            problems.append(f"Brand token không còn trong DOCX: {label}")

    if "Lingual" in header_xml or "image" in header_xml or "imagedata" in header_xml:
        checks.append("Header/logo-watermark structure còn tồn tại.")
    else:
        problems.append("Không phát hiện cấu trúc logo/watermark ở header.")

    if "PAGE" in footer_xml and "txbxContent" in footer_xml:
        checks.append("Footer + centered PAGE badge structure OK.")
    else:
        problems.append("Footer mất PAGE field hoặc page badge structure.")

    return QAResult(ok=not problems, checks=checks, problems=problems)


# ===================== OLD-DOCX IMPORTER =====================
TITLE_RE = re.compile(
    r"(?P<doctype>VOCAB\s+BUILDER)\s+(?P<section>UNIT|MODULE)\s+"
    r"(?P<unit>[A-Za-z0-9._-]+)\s*[:\-]\s*(?P<title>.+)",
    re.IGNORECASE,
)


def _clean_cell_text(text: str) -> str:
    return " ".join((text or "").replace("\r", "\n").split())


def _find_vocab_table(doc: Document):
    for table in doc.tables:
        if not table.rows or len(table.rows[0].cells) < 5:
            continue
        head = [_clean_cell_text(c.text).lower() for c in table.rows[0].cells[:5]]
        if (
            head[0] in HEADER_ALIASES["no"]
            and head[1] in HEADER_ALIASES["word"]
            and head[2] in HEADER_ALIASES["type"]
        ):
            return table
    return None


def extract_old_vocab_docx(data: bytes) -> tuple[dict, str, list[str]]:
    """Extract metadata + CSV from an old VOCAB BUILDER .docx.

    Handles vertically merged No. cells by treating consecutive repeated numbers
    as continuation rows of the same word family.
    """
    doc = Document(io.BytesIO(data))
    notes: list[str] = []

    title_text = ""
    for p in doc.paragraphs:
        t = _clean_cell_text(p.text)
        if "VOCAB BUILDER" in t.upper():
            title_text = t
            break

    metadata = {
        "document_type": "VOCAB BUILDER",
        "section_label": "UNIT",
        "unit": "",
        "heading_unit": "",
        "title": "",
    }
    if title_text:
        match = TITLE_RE.search(title_text)
        if match:
            metadata["document_type"] = " ".join(match.group("doctype").upper().split())
            metadata["section_label"] = match.group("section").upper()
            metadata["unit"] = match.group("unit").strip()
            metadata["heading_unit"] = metadata["unit"].split(".", 1)[0]
            metadata["title"] = match.group("title").strip()
        else:
            notes.append(f"Đọc được title nhưng chưa tách tự động hoàn toàn: {title_text}")
    else:
        notes.append("Không tìm thấy dòng title VOCAB BUILDER; hãy nhập metadata thủ công.")

    table = _find_vocab_table(doc)
    if table is None:
        raise ValueError("Không tìm thấy bảng 5 cột No./Word/Type/Pronunciation/Meaning trong file cũ.")

    out = io.StringIO()
    writer = csv.writer(out, lineterminator="\n")
    writer.writerow(["No.", "Word", "Type", "Pronunciation", "Meaning"])

    previous_no = None
    row_count = 0
    family_count = 0
    for row in table.rows[1:]:
        vals = [_clean_cell_text(c.text) for c in row.cells[:5]]
        while len(vals) < 5:
            vals.append("")
        no, word, word_type, pronunciation, meaning = vals[:5]
        if not any(vals):
            continue
        if not word:
            continue

        # python-docx often returns the merged No. value on every merged row.
        # Consecutive repeats therefore become blank continuation rows.
        out_no = no
        if no and previous_no == no:
            out_no = ""
        elif no:
            previous_no = no
            family_count += 1
        elif previous_no is None:
            notes.append("Có hàng từ vựng xuất hiện trước STT đầu tiên.")

        writer.writerow([out_no, word, word_type, pronunciation, meaning])
        row_count += 1

    notes.append(f"Đã trích {family_count} word families / {row_count} vocabulary rows từ file cũ.")
    return metadata, out.getvalue(), notes


# ===================== STREAMLIT UI =====================
BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "templates" / "lx_vocab_template.docx"
FONT_OPTIONS = [
    "Arial",
    "Montserrat",
    "Orbitron",
    "Monument Extended",
    "Times New Roman",
    "Aptos",
    "Calibri",
    "Verdana",
    "Trebuchet MS",
    "Georgia",
    "Cambria",
]


def _font_index(default: str) -> int:
    try:
        return FONT_OPTIONS.index(default)
    except ValueError:
        return 0


def font_select(label: str, default: str, *, key: str, help: str | None = None):
    return st.selectbox(label, FONT_OPTIONS, index=_font_index(default), key=key, help=help)


def _ensure_defaults() -> None:
    defaults = {
        "source_mode": "Dán CSV / Clipboard",
        "unit": "5.1",
        "heading_unit": "5",
        "section_label": "UNIT",
        "document_type": "VOCAB BUILDER",
        "title": "SPECIAL DAYS",
        "filename": "VOCAB BUILDER UNIT 5.1 - SPECIAL DAYS.docx",
        "raw_data": "",
    }
    for k, v in defaults.items():
        st.session_state.setdefault(k, v)


st.set_page_config(page_title="LingualXplore · Vocab Builder", page_icon="📘", layout="wide")
_ensure_defaults()

st.markdown(
    """
    <style>
    .block-container {padding-top: 1.2rem; padding-bottom: 2rem; max-width: 1320px;}
    [data-testid="stSidebar"] {background: #F7F9FA;}
    .lx-title {font-size: 2rem; font-weight: 800; margin-bottom: .1rem;}
    .lx-sub {color:#52636c; margin-bottom: 1rem;}
    .lx-chip {display:inline-block; padding:.25rem .55rem; border-radius:999px;
              background:#eef4f7; margin-right:.35rem; font-size:.85rem;}
    </style>
    <div class="lx-title">LingualXplore · VOCAB Builder V4</div>
    <div class="lx-sub">Tạo mới hoặc chuyển DOCX cũ → theme mới · Header/Footer linh hoạt · bảng liền mạch.</div>
    """,
    unsafe_allow_html=True,
)

# -------------------- SOURCE / IMPORT FIRST (before widgets bound to metadata keys) --------------------
with st.sidebar:
    st.markdown("### 1 · Nguồn dữ liệu")
    source_mode = st.radio(
        "Chọn cách nhập",
        ["Dán CSV / Clipboard", "Chuyển file Word cũ"],
        key="source_mode",
        label_visibility="collapsed",
    )
    uploaded_old = None
    if source_mode == "Chuyển file Word cũ":
        uploaded_old = st.file_uploader("Upload DOCX phiên bản cũ", type=["docx"])
        if uploaded_old is not None:
            payload = uploaded_old.getvalue()
            digest = hashlib.sha1(payload).hexdigest()
            if st.session_state.get("_old_doc_digest") != digest:
                try:
                    meta, csv_text, import_notes = extract_old_vocab_docx(payload)
                    for key in ("unit", "heading_unit", "section_label", "document_type", "title"):
                        if meta.get(key):
                            st.session_state[key] = meta[key]
                    st.session_state["raw_data"] = csv_text
                    st.session_state["_old_doc_digest"] = digest
                    st.session_state["_import_notes"] = import_notes
                    st.session_state.pop("_import_error", None)
                    stem = Path(uploaded_old.name).stem
                    st.session_state["filename"] = f"{stem} - NEW THEME.docx"
                except Exception as exc:
                    st.session_state["_import_error"] = str(exc)

    st.markdown("### 2 · Thông tin tài liệu")
    st.selectbox("MODULE / UNIT", ["UNIT", "MODULE"], key="section_label")
    st.text_input("Sub-unit / badge", key="unit", help="Ví dụ 5.1 – hiển thị trong ô cam bên trái.")
    st.text_input("Số ở heading", key="heading_unit", help="Ví dụ 5 → UNIT 5: SPECIAL DAYS")
    st.text_input("Document type", key="document_type")
    st.text_input("Title", key="title")
    st.text_input("Tên file tải về", key="filename")

    if st.session_state.get("_import_notes") and source_mode == "Chuyển file Word cũ":
        st.success("Đã đọc file cũ")
        for note in st.session_state["_import_notes"]:
            st.caption("• " + note)
    if st.session_state.get("_import_error") and source_mode == "Chuyển file Word cũ":
        st.error(st.session_state["_import_error"])

# -------------------- TABS --------------------
tab_data, tab_title, tab_table, tab_footer, tab_export = st.tabs(
    ["📥 Dữ liệu", "🎨 Header & Title", "🧩 Bảng", "🦶 Footer", "✅ Xuất file"]
)

with tab_data:
    st.markdown("### Dữ liệu từ vựng")
    if source_mode == "Chuyển file Word cũ":
        st.info("File cũ đã được chuyển về dữ liệu 5 cột bên dưới. Bạn có thể sửa trực tiếp trước khi xuất.")
    else:
        st.info("Dán CSV 5 cột: No., Word, Type, Pronunciation, Meaning. No. trống = cùng word family.")
    raw_data = st.text_area(
        "CSV / Clipboard",
        height=430,
        key="raw_data",
        placeholder='No.,Word,Type,Pronunciation,Meaning\n1,special,"adj, n",,\n,specially,adv,,',
    )
    try:
        preview_doc = parse_vocab_csv(
            raw_data,
            unit=st.session_state.unit,
            heading_unit=st.session_state.heading_unit,
            title=st.session_state.title,
            document_type=st.session_state.document_type,
            section_label=st.session_state.section_label,
        )
        c1, c2 = st.columns(2)
        c1.metric("Word families", preview_doc.family_count)
        c2.metric("Vocabulary rows", preview_doc.word_count)
    except Exception:
        pass

# -------------------- HEADER + TITLE --------------------
HEADER_DEFAULTS = [
    (1, "Supplementary", "Arial", 10.5, "#E7A07F", False, False),
    (1, "Vocabulary", "Arial", 10.5, "#E7A07F", False, False),
    (2, "for", "Arial", 10.5, "#E7A07F", False, False),
    (2, "KET", "Orbitron", 10.5, "#7FA9BC", True, False),
    (2, "Learners", "Orbitron", 10.5, "#7FA9BC", True, False),
    (2, "in", "Arial", 10.5, "#E7A07F", False, False),
    (2, "Orwell", "Orbitron", 10.5, "#7FA9BC", True, False),
    (2, "Classes", "Orbitron", 10.5, "#7FA9BC", True, False),
]

with tab_title:
    st.markdown("### Title block")
    a, b, c = st.columns(3)
    with a:
        with st.container(border=True):
            st.markdown("**Badge 5.1**")
            badge_font = font_select("Font", "Orbitron", key="badge_font")
            badge_size = st.slider("Size", 24, 72, 48, key="badge_size")
            badge_color = st.color_picker("Màu chữ", "#FFFFFF", key="badge_color")
            badge_bold = st.checkbox("Bold", True, key="badge_bold")
    with b:
        with st.container(border=True):
            st.markdown("**VOCAB BUILDER**")
            type_font = font_select("Font", "Monument Extended", key="type_font")
            type_size = st.slider("Size", 12, 32, 18, key="type_size")
            type_color = st.color_picker("Màu chữ", "#BF4E14", key="type_color")
            type_bold = st.checkbox("Bold", False, key="type_bold")
    with c:
        with st.container(border=True):
            st.markdown("**UNIT/MODULE 5: SPECIAL DAYS**")
            heading_font = font_select("Font", "Montserrat", key="heading_font")
            heading_size = st.slider("Size", 18, 40, 27, key="heading_size")
            heading_color = st.color_picker("Màu chữ", "#000000", key="heading_color")
            heading_bold = st.checkbox("Bold", True, key="heading_bold")

    title_clearance = st.slider(
        "Khoảng cách an toàn Header → Title (pt)", 18, 60, 28,
        help="Tăng nếu bạn dùng header 3 dòng hoặc font lớn.",
    )

    st.divider()
    st.markdown("### Header linh hoạt")
    st.caption("Mỗi segment có thể là một từ hoặc một cụm; có thể thêm/bớt và đổi line/style độc lập.")
    ac1, ac2, ac3 = st.columns(3)
    with ac1:
        align1 = st.selectbox("Căn dòng 1", ["left", "center", "right"], index=2)
    with ac2:
        align2 = st.selectbox("Căn dòng 2", ["left", "center", "right"], index=2)
    with ac3:
        align3 = st.selectbox("Căn dòng 3", ["left", "center", "right"], index=2)

    segment_count = st.slider("Số segment header", 1, 16, len(HEADER_DEFAULTS))
    header_values = []
    for i in range(segment_count):
        default = HEADER_DEFAULTS[i] if i < len(HEADER_DEFAULTS) else (2, "", "Arial", 10.5, "#E7A07F", False, False)
        line_d, text_d, font_d, size_d, color_d, bold_d, italic_d = default
        with st.container(border=True):
            cols = st.columns([0.7, 2.2, 1.6, 1.0, 0.9, 0.55, 0.55])
            with cols[0]:
                line = st.selectbox("Line", [1, 2, 3], index=line_d - 1, key=f"h_line_{i}")
            with cols[1]:
                text = st.text_input("Text", text_d, key=f"h_text_{i}")
            with cols[2]:
                font = font_select("Font", font_d, key=f"h_font_{i}")
            with cols[3]:
                size = st.number_input("Size", 6.0, 30.0, size_d, 0.5, key=f"h_size_{i}")
            with cols[4]:
                color = st.color_picker("Color", color_d, key=f"h_color_{i}")
            with cols[5]:
                bold = st.checkbox("B", bold_d, key=f"h_bold_{i}")
            with cols[6]:
                italic = st.checkbox("I", italic_d, key=f"h_italic_{i}")
            header_values.append((line, text, font, size, color, bold, italic))

# -------------------- TABLE --------------------
with tab_table:
    st.markdown("### Table theme")
    st.caption("Giữ row rhythm/cell spacing từ bảng gốc; đổi skin để đồng bộ LingualXplore.")
    c1, c2, c3 = st.columns(3)
    with c1:
        with st.container(border=True):
            st.markdown("**Cột No.**")
            table_no_fill = st.color_picker("Màu nền cột No.", "#BF4E14")
            no_font = font_select("Font STT", "Orbitron", key="no_font")
            no_size = st.number_input("Size STT", 7.0, 16.0, 10.0, 0.5)
            no_color = st.color_picker("Màu STT", "#FFFFFF")
            no_bold = st.checkbox("Bold STT", True)
    with c2:
        with st.container(border=True):
            st.markdown("**Header 4 cột còn lại**")
            table_header_fill = st.color_picker("Màu nền", "#0F4761")
            table_header_font = font_select("Font header", "Arial", key="table_header_font")
            table_header_size = st.number_input("Size header", 8.0, 18.0, 10.5, 0.5)
            table_header_color = st.color_picker("Màu chữ header", "#FFFFFF")
            table_header_bold = st.checkbox("Bold header", True)
    with c3:
        with st.container(border=True):
            st.markdown("**Body & Border**")
            body_font = font_select("Font body", "Times New Roman", key="body_font")
            body_size = st.number_input("Size body", 9.0, 16.0, 12.0, 0.5)
            body_color = st.color_picker("Màu body", "#000000")
            outer_border_color = st.color_picker("Outer border", "#0F4761")
            inner_border_color = st.color_picker("Inner grid", "#B8C8D0")
            outer_border_pt = st.number_input("Outer border (pt)", 0.25, 2.0, 1.0, 0.25)
            inner_border_pt = st.number_input("Inner grid (pt)", 0.25, 1.5, 0.5, 0.25)

# -------------------- FOOTER --------------------
FOOTER_DEFAULTS = [
    ("From", "Arial", 9.0, "#BF4E14", False, False),
    ("Learners", "Orbitron", 9.0, "#0F4761", True, False),
    ("to", "Arial", 9.0, "#BF4E14", False, False),
    ("Explorers", "Orbitron", 9.0, "#0F4761", True, False),
]

with tab_footer:
    st.markdown("### Footer")
    left_col, mid_col, right_col = st.columns([1.35, 0.8, 1.35])

    with left_col:
        with st.container(border=True):
            st.markdown("**Slogan bên trái**")
            slogan_align = st.selectbox("Căn slogan", ["left", "center", "right"], index=0)
            footer_seg_count = st.slider("Số segment slogan", 1, 8, 4)
            footer_segments_values = []
            for i in range(footer_seg_count):
                d = FOOTER_DEFAULTS[i] if i < len(FOOTER_DEFAULTS) else ("", "Arial", 9.0, "#0F4761", False, False)
                text_d, font_d, size_d, color_d, bold_d, italic_d = d
                st.markdown(f"Segment {i+1}")
                cc1, cc2 = st.columns([1.1, 1])
                with cc1:
                    text = st.text_input("Text", text_d, key=f"f_text_{i}")
                    font = font_select("Font", font_d, key=f"f_font_{i}")
                with cc2:
                    size = st.number_input("Size", 6.0, 18.0, size_d, 0.5, key=f"f_size_{i}")
                    color = st.color_picker("Color", color_d, key=f"f_color_{i}")
                    b1, b2 = st.columns(2)
                    with b1:
                        bold = st.checkbox("Bold", bold_d, key=f"f_bold_{i}")
                    with b2:
                        italic = st.checkbox("Italic", italic_d, key=f"f_italic_{i}")
                footer_segments_values.append((text, font, size, color, bold, italic))

    with mid_col:
        with st.container(border=True):
            st.markdown("**Số trang**")
            page_fill = st.color_picker("Màu hình tròn", "#BF4E14")
            page_font = font_select("Font số trang", "Orbitron", key="page_font")
            page_size = st.number_input("Size số trang", 7.0, 16.0, 9.0, 0.5)
            page_color = st.color_picker("Màu số", "#FFFFFF")
            page_bold = st.checkbox("Bold số trang", True)
            page_badge_size = st.slider("Đường kính (pt)", 18, 36, 24)
            st.caption("V4 đặt PAGE field ngay trong hình tròn nên số luôn nằm chính giữa.")

    with right_col:
        with st.container(border=True):
            st.markdown("**Thông tin liên hệ**")
            phone_symbol = st.selectbox("Ký hiệu SĐT", ["☎", "☏", "•"], index=0)
            phone = st.text_input("Số điện thoại", "0345286842")
            email_symbol = st.selectbox("Ký hiệu email", ["✉", "@", "•"], index=0)
            email = st.text_input("Email", "email@yourcenter.com")
            contact_font = font_select("Font contact", "Arial", key="contact_font")
            contact_size = st.number_input("Size contact", 6.0, 14.0, 8.0, 0.5)
            phone_color = st.color_picker("Màu SĐT", "#BF4E14")
            email_color = st.color_picker("Màu email", "#0F4761")
            contact_bold = st.checkbox("Bold contact", True)

# -------------------- EXPORT --------------------
with tab_export:
    st.markdown("### Kiểm tra & xuất file")
    st.markdown(
        '<span class="lx-chip">Montserrat heading</span>'
        '<span class="lx-chip">Orange No. rail</span>'
        '<span class="lx-chip">Orbitron numbering</span>'
        '<span class="lx-chip">Editable footer</span>'
        '<span class="lx-chip">Old DOCX converter</span>',
        unsafe_allow_html=True,
    )
    run_qa = st.checkbox("Chạy Content + Brand QA", value=True)
    generate = st.button("TẠO FILE WORD", type="primary", use_container_width=True)

    if generate:
        if not TEMPLATE_PATH.exists():
            st.error(f"Không tìm thấy template: {TEMPLATE_PATH}")
            st.stop()

        filename = st.session_state.filename.strip() or "VOCAB_BUILDER.docx"
        if not filename.lower().endswith(".docx"):
            filename += ".docx"

        document = parse_vocab_csv(
            st.session_state.raw_data,
            unit=st.session_state.unit,
            heading_unit=st.session_state.heading_unit,
            title=st.session_state.title,
            document_type=st.session_state.document_type,
            section_label=st.session_state.section_label,
        )
        messages = validate_vocab(document)
        for message in messages:
            (st.error if message.level == "error" else st.warning)(message.message)
        if has_errors(messages):
            st.stop()

        header_segments = [
            HeaderSegment(
                int(line), text,
                TextStyle(font, float(size), color.lstrip("#"), bool(bold), bool(italic)),
            )
            for line, text, font, size, color, bold, italic in header_values
            if text.strip()
        ]
        footer_segments = [
            HeaderSegment(
                1, text,
                TextStyle(font, float(size), color.lstrip("#"), bool(bold), bool(italic)),
            )
            for text, font, size, color, bold, italic in footer_segments_values
            if text.strip()
        ]

        theme = ThemeConfig(
            badge=TextStyle(badge_font, badge_size, badge_color.lstrip("#"), badge_bold),
            document_type=TextStyle(type_font, type_size, type_color.lstrip("#"), type_bold),
            heading=TextStyle(heading_font, heading_size, heading_color.lstrip("#"), heading_bold),
            header_segments=header_segments,
            header_alignments={1: align1, 2: align2, 3: align3},
            title_clearance_pt=title_clearance,
            table=TableTheme(
                no_fill=table_no_fill.lstrip("#"),
                header_fill=table_header_fill.lstrip("#"),
                no_header_text=TextStyle(no_font, no_size, no_color.lstrip("#"), no_bold),
                header_text=TextStyle(table_header_font, table_header_size, table_header_color.lstrip("#"), table_header_bold),
                no_body_text=TextStyle(no_font, no_size, no_color.lstrip("#"), no_bold),
                body_text=TextStyle(body_font, body_size, body_color.lstrip("#"), False, False),
                outer_border_color=outer_border_color.lstrip("#"),
                outer_border_pt=outer_border_pt,
                inner_border_color=inner_border_color.lstrip("#"),
                inner_border_pt=inner_border_pt,
            ),
            footer=FooterTheme(
                slogan_segments=footer_segments,
                slogan_align=slogan_align,
                phone=phone,
                email=email,
                phone_symbol=phone_symbol,
                email_symbol=email_symbol,
                contact_font=contact_font,
                contact_size=contact_size,
                phone_color=phone_color.lstrip("#"),
                email_color=email_color.lstrip("#"),
                contact_bold=contact_bold,
                page_fill=page_fill.lstrip("#"),
                page_text=TextStyle(page_font, page_size, page_color.lstrip("#"), page_bold),
                page_badge_size_pt=page_badge_size,
            ),
        )

        st.write(f"Đã nhận **{document.family_count} word families / {document.word_count} từ**.")
        try:
            output = render_vocab_docx(document, template_path=TEMPLATE_PATH, theme=theme)
        except Exception as exc:
            st.exception(exc)
            st.stop()

        if run_qa:
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(output.getvalue())
                tmp_path = Path(tmp.name)
            try:
                result = run_content_and_brand_qa(tmp_path, document)
                if result.ok:
                    st.success("QA semantic + brand structure: PASS")
                else:
                    st.warning("QA có cảnh báo; hãy kiểm tra file Word trực quan.")
                with st.expander("Chi tiết QA"):
                    for item in result.checks:
                        st.write("✅", item)
                    for item in result.problems:
                        st.write("❌", item)
            finally:
                tmp_path.unlink(missing_ok=True)

        st.download_button(
            "TẢI FILE WORD",
            output.getvalue(),
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True,
        )
